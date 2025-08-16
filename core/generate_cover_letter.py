from docx import Document
from io import BytesIO

def create_cover_letter():
    doc = Document()
    doc.add_heading('Cover Letter', 0)
    doc.add_paragraph(
        "Dear Hiring Manager,\n\n"
        "I am excited to apply for the Python Backend Engineer role. With hands-on experience in Django, AWS, "
        "and backend systems, I believe I would make a strong addition to your team.\n\n"
        "Thank you for your time and consideration.\n\nSincerely,\nJohn Doe"
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
