from docx import Document
from io import BytesIO

def create_resume():
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Python Developer\nLeeds, UK\nEmail: johndoe@example.com')

    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Software Engineer at XYZ\n- Built APIs using Django\n- Integrated AWS services')

    doc.add_heading('Education', level=1)
    doc.add_paragraph('BSc Computer Science – University of Bolton')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
